import os
import sys
import pandas as pd
import json
import re
import logging
import datetime
import pdfplumber
import asyncio
import time
import openpyxl
import shutil
from docx import Document
from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter
from openpyxl.chart import LineChart, Reference
from dotenv import load_dotenv
import gc

# Google ADK Imports
from google.adk.agents import LlmAgent, Agent
from google.adk.models.google_llm import Gemini
from google.adk.tools import google_search, AgentTool, preload_memory
from google.adk.sessions import DatabaseSessionService, InMemorySessionService
from google.adk.memory import InMemoryMemoryService
from google.adk.plugins.logging_plugin import LoggingPlugin
from google.genai import types
from google.adk.runners import Runner

# --- 1. GLOBAL SETUP ---
load_dotenv()
APP_NAME = "bioprocess_app"
USER_ID = "user_default"

# --- 2. CONFIGURATION & FOLDERS ---
PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT")
LOCATION = os.environ.get("GOOGLE_CLOUD_LOCATION", "us-central1")
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")

FOLDERS = {
    "INPUT": "input_files",
    "PROCESSED": "processed_files",
    "FINAL": "final_reports"
}
for f in FOLDERS.values():
    os.makedirs(f, exist_ok=True)

logging.basicConfig(level=logging.INFO, force=True)
logger = logging.getLogger(__name__)

# FIX: Aggressive Retry Logic for Error 429 (Rate Limits)
retry_config = types.HttpRetryOptions(
    attempts=15,        # Increased to survive 60s bans
    exp_base=2,         # Exponential backoff
    initial_delay=20,   # Start with 20s wait
    http_status_codes=[429, 500, 503]
)

# --- 3. HELPER FUNCTIONS ---

def resolve_path(file_path):
    """Smart path finder that looks in specific folders."""
    if not file_path: return None
    clean = file_path.strip('"').strip("'").replace('\\', '/')
    
    # 1. Check if exact path exists
    if os.path.exists(clean): return os.path.abspath(clean)
    
    # 2. Check in specific folders (Priority: Processed -> Input)
    basename = os.path.basename(clean)
    for folder in [FOLDERS["PROCESSED"], FOLDERS["INPUT"]]:
        candidate = os.path.join(folder, basename)
        if os.path.exists(candidate): return os.path.abspath(candidate)
    return None

def smart_read_file(f_path):
    real_path = resolve_path(f_path)
    if not real_path: return pd.DataFrame()
    try:
        if real_path.endswith('.xlsx'): return pd.read_excel(real_path)
        try: return pd.read_csv(real_path, sep=None, engine='python')
        except: return pd.read_csv(real_path, sep=';')
    except: return pd.DataFrame()

def profile_file_content(file_path: str):
    try:
        df = smart_read_file(file_path)
        if df.empty: return "Empty file."
        return f"File: {file_path}\nColumns: {list(df.columns)}\nStats:\n{df.describe(include='all').to_string()}"
    except Exception as e: return f"Error: {e}"

async def auto_save_to_memory(callback_context):
    """Automatically saves session to memory."""
    try:
        if hasattr(callback_context, '_invocation_context'):
            ctx = callback_context._invocation_context
            if hasattr(ctx, 'memory_service') and hasattr(ctx, 'session'):
                await ctx.memory_service.add_session_to_memory(ctx.session)
    except: pass

# --- 4. THE TOOLKIT ---

def list_processed_files(directory: str = "processed_files"):
    """
    Lists all files currently waiting in the 'processed_files' folder.
    Use this to find the EXACT filenames before trying to merge or build reports.
    """
    try:
        # Use global FOLDERS dict if available, else default
        target_dir = FOLDERS.get("PROCESSED", "processed_files")
        if not os.path.exists(target_dir):
            return "Error: Directory not found."
            
        files = os.listdir(target_dir)
        if not files:
            return "Folder is empty."
            
        return f"📂 FILES IN '{target_dir}':\n" + "\n".join(files)
    except Exception as e:
        return f"Error listing files: {e}"
        
def inspect_file_headers(file_path: str):
    try:
        df = smart_read_file(file_path)
        if df.empty: return "Error: Empty."
        return f"FILE: {file_path}\nSHAPE: {df.shape}\nCOLUMNS: {list(df.columns)}\nSAMPLE:\n{df.head(3).to_string()}"
    except Exception as e: return f"Error: {e}"

def execute_pandas_operation(file_path: str, operation_description: str, python_code: str):
    """
    Executes dynamic Python code. 
    FIX: Injects 'glob' and 'os' for file finding. Enforces print suppression.
    """
    import io
    import sys
    
    try:
        real_path = resolve_path(file_path)
        # Allow running without a specific file if checking a folder
        if not real_path and "processed_files" not in python_code: 
             return f"Error: File {file_path} not found."
        
        print(f"[SYSTEM] 🐍 Executing Python: {operation_description}")
        
        # Load primary file if it exists
        df = smart_read_file(real_path) if real_path else pd.DataFrame()
        
        import difflib
        import numpy as np
        import glob  
        import os    
        
        # 1. SETUP CAPTURE
        captured_output = io.StringIO()
        original_stdout = sys.stdout
        sys.stdout = captured_output 
        
        # 2. DEFINE SCOPE
        execution_scope = {
            'df': df,
            'pd': pd,
            'np': np,
            'glob': glob,       
            'os': os,           
            'difflib': difflib,
            'datetime': datetime,
            'file_path': real_path,
            'smart_read_file': smart_read_file, 
            'resolve_path': resolve_path
        }
        
        try:
            # 3. RUN CODE
            exec(python_code, execution_scope)
        except Exception as e:
            sys.stdout = original_stdout
            return f"PYTHON ERROR: {e}"
        finally:
            # 4. RESTORE CONSOLE
            sys.stdout = original_stdout
            
        # 5. SAFETY CHECK
        output_str = captured_output.getvalue()
        if len(output_str) > 500: 
            print(f"[SYSTEM] ⚠️ Output truncated: {output_str[:200]}... [omitted]")
        else:
            print(f"[SYSTEM] 🐍 Output: {output_str}")
        
        # 6. SAVE RESULT
        if 'df' in execution_scope:
            df_new = execution_scope['df']
        else:
            # If the code saved files directly (common in merges), that's fine too.
            if "to_csv" in python_code or "to_excel" in python_code:
                return f"SUCCESS: Operation completed. Output checked: {output_str}"
            return "Error: Code did not update 'df' or save a file."
        
        # Auto-save logic
        if real_path:
            base_name = os.path.basename(real_path)
            name_part, ext = os.path.splitext(base_name)
            new_name = f"{name_part}_Processed.csv" if "_Processed" not in name_part else base_name
            output_path = os.path.join(FOLDERS["PROCESSED"], new_name)
            df_new.to_csv(output_path, index=False)
            return f"SUCCESS: Processed {base_name}. Saved to {output_path}. (YOU MUST REPORT THIS)"
        else:
            return f"SUCCESS: General script executed. Output: {output_str}"
        
    except Exception as e: 
        return f"PYTHON ERROR: {e}"

def convert_pdf_to_excel(pdf_path: str, instruction: str = "Extract data"):
    """
    Converts PDF to CSV. 
    FIX: Includes a 'Sleep & Retry' loop to survive Rate Limits (429 Errors).
    """
    try:
        real_path = resolve_path(pdf_path)
        if not real_path: return f"Error: File {pdf_path} not found."
        
        from google import genai
        client = genai.Client(api_key=GOOGLE_API_KEY)
        
        with open(real_path, "rb") as f: pdf_bytes = f.read()
        
        prompt = f"Task: {instruction}. Output raw CSV. No markdown."
        
        max_retries = 5
        wait_seconds = 30 
        
        csv_text = None
        for attempt in range(max_retries):
            try:
                # Using the LITE model as requested
                response = client.models.generate_content(
                    model="gemini-2.5-flash-lite",
                    contents=[types.Content(role="user", parts=[types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"), types.Part.from_text(text=prompt)])]
                )
                csv_text = response.text.replace("```csv", "").replace("```", "")
                break
            except Exception as e:
                # If we hit the rate limit, WAIT instead of crashing
                if "429" in str(e) and attempt < max_retries - 1:
                    print(f"[SYSTEM] ⏳ Rate Limit hit in PDF Tool. Waiting {wait_seconds}s...")
                    time.sleep(wait_seconds)
                    wait_seconds += 10 # Increase wait
                    continue
                else:
                    return f"Error converting PDF: {e}"
        
        if not csv_text: return "Error: Failed to get response."

        base_name = os.path.basename(real_path).replace(".pdf", "_Converted.csv")
        output_path = os.path.join(FOLDERS["PROCESSED"], base_name)
        
        with open(output_path, "w", encoding="utf-8") as f: f.write(csv_text.strip())
        return f"SUCCESS: Converted {base_name}. Saved to {output_path}. (YOU MUST REPORT THIS TO THE USER)"
    except Exception as e: return f"Error: {e}"

def append_offline_sheet(source_file: str, output_filename: str, sheet_name: str = None):
    """Appends data as a sheet. FIX: Handles write/append logic."""
    try:
        real_source = resolve_path(source_file)
        if not real_source: return f"Error: Source {source_file} not found."
        
        output_filename = os.path.basename(output_filename)
        output_path = os.path.join(FOLDERS["FINAL"], output_filename)
            
        df = smart_read_file(real_source)
        
        if os.path.exists(output_path):
            with pd.ExcelWriter(output_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
                df.to_excel(writer, sheet_name=sheet_name or "Data", index=False)
        else:
            with pd.ExcelWriter(output_path, engine='openpyxl', mode='w') as writer:
                df.to_excel(writer, sheet_name=sheet_name or "Data", index=False)
                
        return f"SUCCESS: Appended {sheet_name} to {output_filename}. (YOU MUST REPORT THIS)"
    except Exception as e: return f"Error appending: {e}"

def add_calculated_column(file_path: str, new_col_name: str, formula_template: str):
    try:
        clean_path = resolve_path(file_path)
        if not clean_path: return "Error: File not found."
        wb = load_workbook(clean_path); ws = wb.active 
        col_map = {str(cell.value).strip(): get_column_letter(idx) for idx, cell in enumerate(ws[1], 1) if cell.value}
        new_col_idx = ws.max_column + 1
        ws.cell(row=1, column=new_col_idx, value=new_col_name)
        import re
        ingredients = re.findall(r'\{(.*?)\}', formula_template)
        for row in range(2, ws.max_row + 1):
            row_formula = formula_template
            valid = True
            for ing in ingredients:
                if ing in col_map: row_formula = row_formula.replace(f"{{{ing}}}", f"{col_map[ing]}{row}")
                else: valid = False
            if valid: ws.cell(row=row, column=new_col_idx, value=row_formula)
        wb.save(clean_path)
        return f"SUCCESS: Added {new_col_name}. (YOU MUST REPORT THIS)"
    except Exception as e: return f"Error: {e}"

def add_bioprocess_analysis(file_path: str, graph_column: str = "Dissolved Oxygen"):
    try:
        clean_path = resolve_path(file_path)
        wb = load_workbook(clean_path); ws = wb.active
        headers = [cell.value for cell in ws[1]]
        time_idx = next(i for i, h in enumerate(headers) if h and "Time" in str(h)) + 1
        data_idx = next(i for i, h in enumerate(headers) if h and graph_column in str(h)) + 1
        chart = LineChart()
        chart.title = graph_column; chart.x_axis.title = "Time"
        data = Reference(ws, min_col=data_idx, min_row=1, max_row=ws.max_row)
        cats = Reference(ws, min_col=time_idx, min_row=2, max_row=ws.max_row)
        chart.add_data(data, titles_from_data=True); chart.set_categories(cats)
        ws.add_chart(chart, "E5"); wb.save(clean_path)
        return "SUCCESS: Graph added. (YOU MUST REPORT THIS)"
    except Exception as e: return f"Error: {e}"

def create_formula_reference(filename: str, context: str, equations_dict: str, units_dict: str):
    """
    Creates a professional, human-readable formula sheet with a Units Table.
    """
    try:
        if not filename.endswith(".docx"): filename += ".docx"
        output_path = os.path.join(FOLDERS["FINAL"], filename)
        doc = Document()
        
        # 1. Title & Context
        doc.add_heading('Formula Reference Sheet', 0)
        doc.add_heading('Context', level=1)
        doc.add_paragraph(context)
        
        import ast
        
        # 2. Units Table (Human Readable)
        doc.add_heading('Units & Symbols', level=1)
        try:
            # Safely parse the string input into a dictionary
            units = ast.literal_eval(units_dict) if isinstance(units_dict, str) else units_dict
            
            if isinstance(units, dict):
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Header Row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Symbol'
                hdr_cells[1].text = 'Description'
                hdr_cells[2].text = 'Unit'
                
                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Data Rows
                for symbol, desc_unit in units.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(symbol)
                    
                    # Split "Description (Unit)" if possible
                    if "(" in desc_unit and desc_unit.endswith(")"):
                        desc, unit = desc_unit.rsplit("(", 1)
                        row_cells[1].text = desc.strip()
                        row_cells[2].text = unit.strip(")")
                    else:
                        row_cells[1].text = str(desc_unit)
                        row_cells[2].text = "-"
            else:
                doc.add_paragraph(str(units_dict))
        except Exception as e:
            doc.add_paragraph(f"Could not format table: {units_dict}")

        # 3. Formulas (Clean List)
        doc.add_heading('Formulas', level=1)
        try:
            eqs = ast.literal_eval(equations_dict) if isinstance(equations_dict, str) else equations_dict
            if isinstance(eqs, dict):
                for name, formula in eqs.items():
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    runner = p.add_run(f"{name}: ")
                    runner.bold = True
                    p.add_run(f"\n{formula}")
            else:
                doc.add_paragraph(str(equations_dict))
        except:
            doc.add_paragraph(str(equations_dict))

        doc.save(output_path)
        return f"SUCCESS: Created human-readable report at {output_path}."
    except Exception as e: return f"Error: {e}"

def create_word_report(summary: str, filename: str = "Report.docx"):
    try:
        output_path = os.path.join(FOLDERS["FINAL"], filename)
        doc = Document(); doc.add_heading('Final Report', 0); doc.add_paragraph(summary)
        doc.save(output_path)
        return f"SUCCESS: Report saved to {output_path}. (YOU MUST REPORT THIS)"
    except Exception as e: return f"Error: {e}"

# --- 5. AGENT DEFINITIONS ---

# 1. The Data Wrangler
data_wrangler_agent = LlmAgent(
    name="Data_Wrangler_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Cleans data.",
    instruction="""
    You are a Senior Data Scientist.
    
    ### 🛑 SYNTAX & PATH SAFETY (CRITICAL):
    1. **NO BACKSLASHES:** Windows paths (e.g., `C:\\Users...`) cause syntax errors. 
       - **ALWAYS** use relative paths: `processed_files/MyFile.csv`.
    2. **USE GLOB:** To merge multiple files, use:
       `files = glob.glob('processed_files/*_Converted.csv')`
       Then loop through `files` to read and concat. Do NOT hardcode file paths.
    3. **STATELESS:** Variables (`df1`, `df2`) DO NOT persist between tool calls. You must re-read files to merge them.
    
    ### 🛑 INPUT STRICTNESS:
    1. **NO GUESSING:** Do NOT invent filenames. Use the exact paths provided or find them with `glob`.
    2. **FILTERING:** If the user mentions "R06", it refers to DATA INSIDE the file, not the filename.
    
    ### 🔧 AUTO-CORRECTION STRATEGY:
    If you encounter "Column mismatch" or "Length mismatch" errors during merging:
    1. **DO NOT GIVE UP.**
    2. **INSPECT:** Use `df.columns` to see the *actual* column names.
    3. **STANDARDIZE:** Rename columns to a common standard before merging.
       - *Example:* Rename 'Sample', 'Sample Name', 'Name' -> 'Sample_ID'.
    4. **RETRY:** Attempt the merge again with the corrected columns.
    
    ### ⚡️ EXECUTION RULES:
    1. **PROCESS ALL FILES:** Process every file in the list.
    2. **MERGING:** Use the injected `smart_read_file(f)` inside your loop.
    3. **NO PRINTING:** Do NOT print entire dataframes. Only print `.head()`.
    4. **RENAMING:** Output files MUST use a different name (e.g. `_Processed.csv`).
    
    ### 🗣️ MANDATORY REPORTING (ANTI-CRASH RULE):
    - **YOU MUST SPEAK:** After running tools, you MUST return a text summary.
    - **FORBIDDEN:** Do NOT return empty content.
    - **TEMPLATE:** "I have successfully processed: [LIST FILES]. They are saved in `processed_files/`."
    """,
    tools=[inspect_file_headers, execute_pandas_operation, convert_pdf_to_excel]
)

# 2. The Excel Architect
excel_architect_agent = LlmAgent(
    name="Excel_Architect_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Builds Excel.",
    instruction="""
    You are the Master Builder.
    
    ### 👁️ REALITY CHECK (CRITICAL):
    1. **FIRST ACTION:** You MUST call `list_processed_files` immediately to see what files actually exist in `processed_files/`.
    2. **NEVER GUESS:** Do not assume filenames like "Online_Data.csv" or "table_1.csv". 
    3. **USE FACTUAL NAMES:** Only use the exact filenames returned by the list tool.
    
    ### 🛠️ CAPABILITIES:
    - Your tool `append_offline_sheet` SUPPORTS both .csv and .xlsx input files.
    - **NEVER** complain about file formats. Just try to append them.
    
    ### 🛠️ TASK:
    1. List the files.
    2. Use `append_offline_sheet` to build 'Consolidated_Data.xlsx' in `final_reports/` using the actual files you found.
    
    ### 🏁 REPORTING:
    - Return a text summary of exactly which sheets were added.
    """,
    tools=[append_offline_sheet, list_processed_files] # <--- Added the new tool here
)

# Search Agent
web_search_agent = LlmAgent(
    name="Web_Search_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Searches Google.",
    instruction="""
    You are a Web Researcher.
    - Your ONLY job is to search Google for information.
    - Summarize the search results clearly for the user.
    """,
    tools=[google_search]
)

# 4. The Research Agent (UPDATED: Human-Readable Formatting)
research_agent = LlmAgent(
    name="Research_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Formula Expert.",
    instruction="""
    You are a Scientific Researcher. 
    - **Job:** Create clear, human-readable documentation for the user.
    - **Tool Use:** When using `create_formula_reference`, you MUST provide a `units_dict`.
    - **Format Rule:** The `units_dict` values MUST follow the format: "Description (Unit)".
      - *Good:* `{"t": "Time (hours)", "C_glc": "Glucose Conc. (g/L)"}`
      - *Bad:* `{"t": "Time"}`
    - **Formulas:** Use standard mathematical notation where possible in `equations_dict`.
    """,
    tools=[create_formula_reference] 
)

# 5. Excel Formula Agent (UPDATED: Anti-Silence)
formula_agent = LlmAgent(
    name="Excel_Formula_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Adds Excel formulas.",
    instruction="""
    You are an Excel Expert.
    
    ### 🛠️ TASK:
    - Inject formulas into 'Consolidated_Data.xlsx' using `add_calculated_column`.
    - If you need column names, use `inspect_file_headers` first.
    
    ### 🗣️ MANDATORY REPORTING:
    - **CRITICAL:** You must NEVER return an empty response.
    - **CRITICAL:** After adding columns, you MUST write a text summary.
    - **Example:** "I have successfully added the following formula columns: [List Columns]."
    - If you are silent, the system will crash. **ALWAYS SPEAK.**
    """,
    tools=[add_calculated_column, inspect_file_headers]
)

excel_analyst_agent = LlmAgent(
    name="Excel_Analyst_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Creates graphs.",
    instruction="Ask user what to graph, then use `add_bioprocess_analysis`.",
    tools=[add_bioprocess_analysis, profile_file_content]
)

word_report_agent = LlmAgent(
    name="Word_Report_Agent",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Writes report.",
    instruction="Summarize findings into 'Report.docx' in the FINAL folder.",
    tools=[create_word_report]
)

# --- 6. ROOT AGENT (The Project Manager) ---

Bioprocess_agent = LlmAgent(
    name="Bioprocess_Manager",
    model=Gemini(model="gemini-2.5-flash-lite", api_key=GOOGLE_API_KEY, retry_options=retry_config),
    description="Manager.",
    instruction="""
    You are the Project Manager.
    
    ### 🗣️ TRANSPARENCY PROTOCOL (MANDATORY):
    - **EXPLAIN FIRST:** Before calling *any* tool, you must write a short sentence explaining your plan to the user.
    - **Format:** "I am now [Action] because [Reason]..."
    - **Example:** "I am calling the Data Wrangler to clean the raw PDF files so we can use them in Excel."
    
    ### 🚨 FILE ACCESS PROTOCOL (CRITICAL)
    1. **TRUST THE PATHS:** The user has provided local file paths in `[SYSTEM_DATA]`.
    2. **YOU HAVE PERMISSION:** Do not verify access. Just pass the strings to the tools.
    3. **DELEGATE:** Your job is to pass these exact paths to `Data_Wrangler_Agent`.
    
    ### 🛠️ TOOL CONTRACT (INPUTS & OUTPUTS)
    
    **1. Data_Wrangler_Agent**
    - **INPUT:** Raw file paths from `[SYSTEM_DATA]`.
    - **OUTPUT:** Cleaned **CSV files** in `processed_files/`.
    - **LIMITATION:** This agent CANNOT create the final Excel file. It only makes intermediate CSVs.
    
    **2. Excel_Architect_Agent**
    - **INPUT:** Explicit list of filenames in `processed_files/`.
    - **OUTPUT:** A single **.xlsx file** in `final_reports/`.
    - **MANDATE:** You MUST pass the EXACT filenames created by the Wrangler to this agent.
    
    **3. Web_Search_Agent**
    - **INPUT:** A specific question.
    - **OUTPUT:** Search results from Google.
    - **USE CASE:** Only for general questions (e.g., "What is glucose?").
    
    ### 🧠 CORE BEHAVIOR: DECIDE -> PLAN -> EXECUTE
    
    **SCENARIO A: RAW FILES (PDFs/CSVs)**
    *Trigger:* User uploads raw data files.
    
    1. **PROPOSE PLAN:** Read `[SYSTEM_DATA]`. Output a numbered plan (Extraction -> Assembly -> Formulas).
    2. **ASK:** "Does this plan look correct? Say 'Yes' to execute."
    3. **EXECUTE (After 'Yes'):** Run the pipeline below without stopping until Phase 3.
    
       - **PHASE 1: WRANGLER (Extraction)**
         - *Explain:* "Starting extraction of raw files..."
         - Call `Data_Wrangler_Agent`. Pass the **EXACT** file paths list.
         - *Wait for confirmation that CSVs are saved.*
         
       - **PHASE 1.5: VERIFICATION (Reality Check)**
         - *Explain:* "Verifying the generated files..."
         - **ACTION:** Call `list_processed_files` to see exactly what files exist.
         - *Wait for the list of filenames.*
         
       - **PHASE 2: ARCHITECT (Assembly)**
         - *Explain:* "Assembling the Excel report from verified CSVs..."
         - Call `Excel_Architect_Agent`.
         - **CRITICAL:** Use the filenames found in Phase 1.5.
         - *Command:* "Build `final_reports/Consolidated_Data.xlsx` using [INSERT EXACT FILENAMES FROM LIST]."
         
       - **PHASE 3: RESEARCH (Formulas)**
         - *Explain:* "Checking required scientific formulas..."
         - Call `Research_Agent`.
         - **STOP:** Ask user to verify formulas.
         
       - **PHASE 4: COMPLETION**
         - Call `Excel_Formula_Agent` -> `Excel_Analyst_Agent` -> `Word_Report_Agent`.
    
    **SCENARIO B: CLEAN EXCEL**
    *Trigger:* User uploads a file named 'Chemostat_Processed_Data.xlsx'.
    1. **ACTION:** Skip directly to **Phase 3** (Research/Formulas).
    
    **SCENARIO C: GENERAL QUESTIONS**
    *Trigger:* User asks a question (e.g., "What is a chemostat?") without files.
    1. **ACTION:** Call `Web_Search_Agent` immediately.
    
    ### 🏁 STATUS UPDATE:
    **Project Status:**
    * 📊 Consolidated Data: [Pending/Done]
    * 📝 Formula Sheet: [Pending/Done]
    * 📑 Final Report: [Pending/Done]
    """,
    tools=[
        preload_memory, 
        list_processed_files,
        AgentTool(agent=data_wrangler_agent), 
        AgentTool(agent=excel_architect_agent), 
        AgentTool(agent=research_agent), 
        AgentTool(agent=web_search_agent),
        AgentTool(agent=formula_agent), 
        AgentTool(agent=excel_analyst_agent), 
        AgentTool(agent=word_report_agent)
    ],
    after_agent_callback=auto_save_to_memory
)

# --- 7. RUNNER ---
session_service = InMemorySessionService()
memory_service = InMemoryMemoryService()
logging_plugin = LoggingPlugin()

auto_runner = Runner(
    agent=Bioprocess_agent, 
    app_name=APP_NAME, 
    session_service=session_service, 
    memory_service=memory_service, 
    plugins=[logging_plugin]
)